from docx import Document
import subprocess
import os
import questionary

# Criar um arquivo Word
doc = Document()

# Adicionar informações ao documento
doc.add_paragraph("CHECKLIST ✔")
doc.add_paragraph("PROJETOS DE IRRIGAÇÃO")

# Questionário interativo
nome_obra_projeto_cliente = questionary.text("NOME/OBRA/PROJETO/CLIENTE:").ask()
prazo_entrega = questionary.text("Prazo estimado para entrega (ex. 10/07/2023):").ask()
area_paisagismo_torre_a = questionary.text("Área total de paisagismo - Torre A (Norte):").ask()
area_paisagismo_torre_b = questionary.text("Área total de paisagismo - Torre B (Sul):").ask()
area_paisagismo_lazer = questionary.text("Área total de paisagismo - Lazer:").ask()
origem_agua_irrigacao = questionary.text("Origem/tipo/tamanho da água para irrigação:").ask()
distancia_reservatorio_ultimo_pavimento = questionary.text("Distância do reservatório até último pavimento irrigado:").ask()
regulagem_pressao_pavimento = questionary.text("Regulagem de pressão pavimento:").ask()
localizacao_area_tecnica_sistema_a = questionary.text("Localização da área técnica - Sistema A:").ask()
localizacao_area_tecnica_sistema_b = questionary.text("Localização da área técnica - Sistema B:").ask()
localizacao_area_tecnica_sistema_c = questionary.text("Localização da área técnica - Sistema C:").ask()
num_pavimentos = questionary.text("Nº pavimentos:").ask()
com_paisagismo = questionary.text("Com paisagismo:").ask()

# Adicionar respostas ao documento
info = f"""
NOME/OBRA/PROJETO/CLIENTE: {nome_obra_projeto_cliente}
Prazo estimado para entrega: {prazo_entrega}
Área total de paisagismo: 
Torre A (Norte) = {area_paisagismo_torre_a} m²
Torre B (Sul) = {area_paisagismo_torre_b} m²
Lazer {area_paisagismo_lazer} m²
Origem/tipo/tamanho da água para irrigação: {origem_agua_irrigacao}
Distância do reservatório até último pavimento irrigado: {distancia_reservatorio_ultimo_pavimento} mca
Regulagem de pressão pavimento: {regulagem_pressao_pavimento}
Localização da área técnica:
Sistema A: {localizacao_area_tecnica_sistema_a}
Sistema B: {localizacao_area_tecnica_sistema_b}
Sistema C: {localizacao_area_tecnica_sistema_c}
Nº pavimentos: {num_pavimentos}
Com paisagismo: {com_paisagismo}
"""

doc.add_paragraph(info)

# Salvar o arquivo Word
doc.save("meuarquivo.docx")

# Abrir o arquivo Word
if os.name == 'nt':  # Verificar se está rodando no Windows
    subprocess.Popen(["start", "meuarquivo.docx"], shell=True)
elif os.name == 'posix':  # Verificar se está rodando no Linux
    subprocess.Popen(["xdg-open", "meuarquivo.docx"])
else:
    print("Sistema operacional não suportado para abrir automaticamente o Word.")
